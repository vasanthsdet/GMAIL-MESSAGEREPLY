#!/usr/bin/env python3
"""
Gmail Resume Workflow
---------------------
1. Authenticate with Gmail API (OAuth2)
2. Fetch today's emails
3. Detect resume requests + extract required skills via Claude Haiku
4. Update resume DOCX (add skills to 2 project sections, preserve formatting)
5. Reply to the email with the updated resume attached
"""

import truststore; truststore.inject_into_ssl()

import os
import base64
import pickle
import json
import re
from datetime import date
from pathlib import Path
import json
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from docx import Document
from docx.oxml.ns import qn
import anthropic

# ─── Paths & Config ───────────────────────────────────────────────────────────

BASE_DIR         = Path(__file__).parent
CREDENTIALS_FILE = BASE_DIR / "credentials.json"
TOKEN_FILE       = BASE_DIR / "token.pickle"

# Place your resume in the same folder as this script; change the name here if needed.
RESUME_FILE         = BASE_DIR / "Revathi Battina- Resume.docx"
UPDATED_RESUME_FILE = BASE_DIR / "resume_updated.docx"
PROCESSED_IDS_FILE   = BASE_DIR / "processed_ids.json"
REPLIED_SENDERS_FILE = BASE_DIR / "replied_senders.json"

SCOPES = [
    "https://www.googleapis.com/auth/gmail.readonly",
    "https://www.googleapis.com/auth/gmail.send",
]

# ─── Gmail OAuth2 ─────────────────────────────────────────────────────────────

def get_gmail_service():
    """Authenticate and return an authorised Gmail API service object."""
    creds = None

    if TOKEN_FILE.exists():
        with open(TOKEN_FILE, "rb") as fh:
            creds = pickle.load(fh)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            if not CREDENTIALS_FILE.exists():
                raise FileNotFoundError(
                    f"credentials.json not found at {CREDENTIALS_FILE}\n"
                    "Download it from Google Cloud Console → APIs & Services → Credentials."
                )
            flow = InstalledAppFlow.from_client_secrets_file(
                str(CREDENTIALS_FILE), SCOPES
            )
            creds = flow.run_local_server(port=0, open_browser=False)

        with open(TOKEN_FILE, "wb") as fh:
            pickle.dump(creds, fh)

    return build("gmail", "v1", credentials=creds)

# ─── Processed IDs (deduplication) ───────────────────────────────────────────

def load_processed_ids() -> set:
    if PROCESSED_IDS_FILE.exists():
        return set(json.loads(PROCESSED_IDS_FILE.read_text()))
    return set()

def save_processed_id(msg_id: str) -> None:
    ids = load_processed_ids()
    ids.add(msg_id)
    PROCESSED_IDS_FILE.write_text(json.dumps(list(ids)))

def load_replied_senders() -> set:
    if REPLIED_SENDERS_FILE.exists():
        return set(json.loads(REPLIED_SENDERS_FILE.read_text()))
    return set()

def save_replied_sender(sender: str) -> None:
    senders = load_replied_senders()
    senders.add(_extract_email(sender))
    REPLIED_SENDERS_FILE.write_text(json.dumps(list(senders)))

def _extract_email(address: str) -> str:
    """Pull bare email from 'Name <email@x.com>' or return as-is."""
    match = re.search(r"<([^>]+)>", address)
    return match.group(1).lower() if match else address.lower()

# ─── Fetch Latest Emails ──────────────────────────────────────────────────────

def get_latest_emails(service, max_results: int = 30) -> list[dict]:
    """Return the latest inbox messages, skipping already-processed ones."""
    processed = load_processed_ids()
    try:
        result   = service.users().messages().list(
            userId="me", q="in:inbox", maxResults=max_results
        ).execute()
        messages = result.get("messages", [])
    except HttpError as exc:
        print(f"[Gmail] Error listing messages: {exc}")
        return []

    emails = []
    for msg in messages:
        if msg["id"] in processed:
            continue
        try:
            full = service.users().messages().get(
                userId="me", id=msg["id"], format="full"
            ).execute()
            emails.append(_parse_message(full))
        except HttpError as exc:
            print(f"[Gmail] Error fetching message {msg['id']}: {exc}")

    return emails


def _parse_message(message: dict) -> dict:
    headers = {h["name"]: h["value"] for h in message["payload"]["headers"]}
    return {
        "id":        message["id"],
        "thread_id": message["threadId"],
        "from":      headers.get("From", ""),
        "to":        headers.get("To", ""),
        "subject":   headers.get("Subject", "(no subject)"),
        "date":      headers.get("Date", ""),
        "body":      _extract_text(message["payload"]),
        "snippet":   message.get("snippet", ""),
    }


def _extract_text(payload: dict) -> str:
    """Recursively extract the best plain-text representation of a part."""
    mime = payload.get("mimeType", "")

    if mime == "text/plain":
        data = payload.get("body", {}).get("data", "")
        if data:
            return base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")

    if mime == "text/html":
        data = payload.get("body", {}).get("data", "")
        if data:
            html = base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")
            return re.sub(r"<[^>]+>", " ", html)

    for part in payload.get("parts", []):
        text = _extract_text(part)
        if text.strip():
            return text

    return ""

# ─── Skill Detection via Claude Haiku ────────────────────────────────────────

def detect_resume_request(email: dict) -> dict:
    """
    Ask Claude Haiku whether the email requests a resume with specific skills.

    Returns:
        {
          "is_resume_request": bool,
          "skills": list[str],    # concrete technical skills mentioned
          "reason": str
        }
    """
    client = anthropic.Anthropic()   # reads ANTHROPIC_API_KEY from env

    email_text = (
        f"Subject: {email['subject']}\n"
        f"From: {email['from']}\n\n"
        f"{email['body'][:3000]}"    # cap tokens
    )

    already_replied = _extract_email(email["from"]) in load_replied_senders()

    prompt = f"""Analyze the email below and determine whether it is requesting a resume or CV that highlights specific technical skills.

Email:
---
{email_text}
---

Reply with ONLY a valid JSON object — no other text:
{{
  "is_resume_request": true or false,
  "is_update_request": true or false,
  "skills": ["Skill1", "Skill2"],
  "reason": "one-line explanation"
}}

Rules:
- Set is_resume_request to true if the sender asks for a resume/CV (explicitly or implicitly, e.g. "please share your updated resume", "we need someone with X, please apply").
- Set is_update_request to true ONLY if the sender is explicitly asking for a NEW or UPDATED resume with different/additional skills compared to what was sent before (e.g. "can you send an updated resume highlighting X", "we now need Y skills", "please resend with focus on Z").
- List every concrete technical skill, language, framework, tool, or platform mentioned (e.g. "Python", "React", "AWS", "Docker").
- Exclude soft skills ("communication", "leadership") and vague terms ("experience").
- If no skills are mentioned, return an empty array.
- Return false for is_resume_request if it is a newsletter, promotion, follow-up, thank-you, or unrelated message."""

    response = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=512,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = response.content[0].text.strip()
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if match:
        try:
            result = json.loads(match.group())
            result["_already_replied"] = already_replied
            return result
        except json.JSONDecodeError:
            pass

    print(f"[Claude] Unexpected response: {raw[:200]}")
    return {"is_resume_request": False, "is_update_request": False, "skills": [], "reason": "parse error", "_already_replied": already_replied}

# ─── Human-Written Text via Claude ───────────────────────────────────────────


def _write_reply_text(original_email: dict, skills: list[str]) -> str:
    """
    Ask Claude Haiku to write a brief, natural reply email — not a template.
    Matches the tone of the recruiter's email.
    """
    client = anthropic.Anthropic()
    prompt = f"""A recruiter sent me this email:

Subject: {original_email['subject']}
{original_email['body'][:700]}

Write a short, genuine reply from me (the job seeker). I am attaching my updated resume that now highlights: {", ".join(skills)}.

Requirements:
- 2–3 sentences only
- Sound like a real person, not a template
- Reference the specific skills naturally
- Match the recruiter's tone (formal if they are formal, casual if casual)
- Do NOT use filler phrases like "Hope this email finds you well", "Please don't hesitate to reach out", "I look forward to hearing from you", or "Best regards"
- End the reply naturally without a sign-off line

Output only the email body text."""

    response = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=200,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip()


# ─── DOCX Resume Editing ──────────────────────────────────────────────────────

def _find_experience_sections(doc: Document) -> list[dict]:
    """
    Detect job/experience sections by finding paragraphs that contain a date range.
    Returns a list of dicts sorted by appearance (most recent first in typical resumes).
    """
    date_rx = re.compile(
        r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b.{0,10}\d{4}',
        re.IGNORECASE,
    )
    paras    = doc.paragraphs
    sections = []
    pending  = None   # (start_idx, title)

    for i, para in enumerate(paras):
        text = para.text.strip()
        if not text:
            continue
        if date_rx.search(text):
            if pending:
                sections.append({
                    "start_idx": pending[0],
                    "end_idx":   i - 1,
                    "title":     pending[1],
                })
            pending = (i, text)

    if pending:
        sections.append({
            "start_idx": pending[0],
            "end_idx":   len(paras) - 1,
            "title":     pending[1],
        })

    return sections


def _last_content_para(doc: Document, start: int, end: int):
    """Return the last non-empty paragraph within [start, end]."""
    paras = doc.paragraphs
    for i in range(min(end, len(paras) - 1), start - 1, -1):
        if paras[i].text.strip():
            return paras[i]
    return None


def _generate_skill_bullet(role_title: str, skills: list[str]) -> str:
    """Ask Claude Haiku to write one human-sounding bullet point about the skills."""
    client = anthropic.Anthropic()
    prompt = f"""You are editing a QA Automation Engineer's resume. Write ONE bullet point sentence describing how they used these technologies in their role.

Role context: {role_title}
Skills/technologies to mention: {", ".join(skills)}

Rules:
- Exactly one sentence
- Start with a strong past-tense action verb (Leveraged, Applied, Utilized, Implemented, Integrated, etc.)
- Sound like a real person wrote it — not AI-generated
- Be concrete and professional (QA/engineering context)
- Do NOT end with a period (this resume style omits periods)
- Do NOT add bullet symbols, hyphens, or quotes
- Output the sentence text only"""

    response = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=120,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.content[0].text.strip().lstrip("•–-· ")


def _insert_paragraph_after(ref_para, text: str) -> None:
    """
    Insert a new paragraph immediately after `ref_para`, cloning its XML
    structure so fonts, spacing, and indentation match exactly.
    """
    import copy
    from lxml import etree

    new_elem = copy.deepcopy(ref_para._element)

    # Strip all runs and hyperlinks from the clone
    for tag in (qn("w:r"), qn("w:hyperlink"), qn("w:ins"), qn("w:del")):
        for node in new_elem.findall(tag):
            new_elem.remove(node)

    # Build a new run, copying rPr from the last run of ref_para
    r_elem = etree.SubElement(new_elem, qn("w:r"))
    last_runs = ref_para._element.findall(qn("w:r"))
    if last_runs:
        rpr = last_runs[-1].find(qn("w:rPr"))
        if rpr is not None:
            r_elem.insert(0, copy.deepcopy(rpr))

    t_elem = etree.SubElement(r_elem, qn("w:t"))
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    ref_para._element.addnext(new_elem)


def update_resume(skills: list[str]) -> bool:
    """
    Insert a Claude-written human bullet point about `skills` into the
    2 most recent experience sections. All existing formatting is preserved.
    """
    if not RESUME_FILE.exists():
        print(f"[DOCX] Resume not found: {RESUME_FILE}")
        return False
    if not skills:
        print("[DOCX] No skills to add.")
        return False

    doc      = Document(str(RESUME_FILE))
    sections = _find_experience_sections(doc)

    if not sections:
        print("[DOCX] No experience sections detected in resume.")
        return False

    added = 0
    for section in sections[:2]:
        title     = section["title"]
        last_para = _last_content_para(doc, section["start_idx"], section["end_idx"])
        if last_para is None:
            continue

        print(f"[Claude] Writing bullet for → {title[:65]} …")
        bullet = _generate_skill_bullet(title, skills)
        print(f'[DOCX]   + "{bullet}"')
        _insert_paragraph_after(last_para, bullet)
        added += 1

    if added == 0:
        print("[DOCX] No sections updated.")
        return False

    doc.save(str(UPDATED_RESUME_FILE))
    print(f"[DOCX] Saved → {UPDATED_RESUME_FILE}")

    # Verify every requested skill appears in the saved file
    saved = Document(str(UPDATED_RESUME_FILE))
    full_text = "\n".join(p.text for p in saved.paragraphs)
    found    = [s for s in skills if s.lower() in full_text.lower()]
    missing  = [s for s in skills if s.lower() not in full_text.lower()]
    print(f"[VERIFY] Skills confirmed in resume : {found}")
    if missing:
        print(f"[VERIFY] Skills NOT found in resume: {missing}")

    return True

# ─── Send Gmail Reply with Attachment ─────────────────────────────────────────

def send_reply(service, original: dict, skills: list[str]) -> bool:
    """Reply to `original` with the updated resume as an attachment."""
    if not UPDATED_RESUME_FILE.exists():
        print("[Gmail] Updated resume not found — cannot send reply.")
        return False

    print("[Claude] Writing reply …")
    body_text = _write_reply_text(original, skills)

    msg = MIMEMultipart()
    msg["To"]         = original["from"]
    msg["Subject"]    = f"Re: {original['subject']}"
    msg["In-Reply-To"] = original["id"]
    msg["References"]  = original["id"]
    msg.attach(MIMEText(body_text, "plain"))

    with open(UPDATED_RESUME_FILE, "rb") as fh:
        part = MIMEBase(
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        part.set_payload(fh.read())

    encoders.encode_base64(part)
    part.add_header("Content-Disposition", 'attachment; filename="Resume.docx"')
    msg.attach(part)

    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8")
    try:
        service.users().messages().send(
            userId="me",
            body={"raw": raw, "threadId": original["thread_id"]},
        ).execute()
        print(f"[Gmail] Reply sent to {original['from']}")
        return True
    except HttpError as exc:
        print(f"[Gmail] Failed to send reply: {exc}")
        return False

# ─── Main Workflow ────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Gmail Resume Workflow")
    print("=" * 60)

    # Step 1 — Authenticate
    print("\n[1/4] Authenticating with Gmail …")
    service = get_gmail_service()
    print("      Done.")

    # Step 2 — Fetch latest 30 unprocessed emails
    print("\n[2/4] Fetching latest 30 emails …")
    emails = get_latest_emails(service, max_results=30)
    print(f"      Found {len(emails)} new email(s).")

    if not emails:
        print("\nNo new emails to process. Exiting.")
        return

    # Step 3 — Detect resume requests
    print("\n[3/4] Scanning emails for resume requests …")
    requests_found = []

    for i, em in enumerate(emails, 1):
        subject = em["subject"][:55]
        sender  = em["from"]
        print(f"  [{i}/{len(emails)}] {subject}")

        if re.search(r"no.?reply", sender, re.IGNORECASE):
            print(f"         → skipped (noreply sender: {sender[:60]})")
            continue

        result = detect_resume_request(em)

        if result.get("is_resume_request") and result.get("skills"):
            if result.get("_already_replied") and not result.get("is_update_request"):
                print(f"         → skipped (already replied to this sender, no update requested)")
            else:
                tag = "UPDATE REQUEST" if result.get("is_update_request") else "RESUME REQUEST"
                print(f"         → {tag}  skills: {result['skills']}")
                requests_found.append({"email": em, "skills": result["skills"]})
        else:
            reason = result.get("reason", "—")
            print(f"         → not a request  ({reason})")

    if not requests_found:
        print("\nNo resume requests found. Exiting.")
        # Mark all scanned emails as processed so we don't re-scan them
        for em in emails:
            save_processed_id(em["id"])
        return

    print(f"\n  {len(requests_found)} resume request(s) to process.")

    # Step 4 — Process each request
    print("\n[4/4] Updating resume and sending replies …")

    for req in requests_found:
        em     = req["email"]
        skills = req["skills"]
        print(f"\n  Email : {em['subject'][:55]}")
        print(f"  Skills: {skills}")

        if update_resume(skills):
            if send_reply(service, em, skills):
                save_replied_sender(em["from"])
        else:
            print("  Skipping reply because resume could not be updated.")
        save_processed_id(em["id"])

    # Mark non-request emails as processed too
    request_ids = {req["email"]["id"] for req in requests_found}
    for em in emails:
        if em["id"] not in request_ids:
            save_processed_id(em["id"])

    print("\n" + "=" * 60)
    print("  Workflow complete.")
    print("=" * 60)


if __name__ == "__main__":
    main()
