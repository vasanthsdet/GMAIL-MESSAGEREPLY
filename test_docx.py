"""Test: Gmail email scan (first 10) + DOCX section detection + Claude bullet generation."""
import os, sys
import truststore; truststore.inject_into_ssl()
sys.path.insert(0, os.path.dirname(__file__))

from main import (
    _find_experience_sections, _generate_skill_bullet, update_resume,
    get_gmail_service, detect_resume_request,
)
from googleapiclient.errors import HttpError
from docx import Document

RESUME = "Revathi Battina- Resume.docx"

# 0. Gmail: fetch and scan first 10 emails
print("=" * 60)
print("Gmail: scanning first 10 emails for resume requests ...")
print("=" * 60)
try:
    service = get_gmail_service()
    result  = service.users().messages().list(userId="me", maxResults=10).execute()
    msgs    = result.get("messages", [])
    print(f"Fetched {len(msgs)} email(s)\n")

    resume_requests = []
    for i, m in enumerate(msgs, 1):
        try:
            full = service.users().messages().get(
                userId="me", id=m["id"], format="full"
            ).execute()
            headers = {h["name"]: h["value"] for h in full["payload"]["headers"]}
            subject = headers.get("Subject", "(no subject)")
            sender  = headers.get("From", "")
            print(f"  [{i:02d}] {subject[:55]}")
            print(f"        From: {sender[:55]}")

            # build minimal email dict for detect_resume_request
            import base64, re
            def _extract(payload):
                mime = payload.get("mimeType", "")
                if mime == "text/plain":
                    data = payload.get("body", {}).get("data", "")
                    if data:
                        return base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")
                if mime == "text/html":
                    data = payload.get("body", {}).get("data", "")
                    if data:
                        html = base64.urlsafe_b64decode(data).decode("utf-8", errors="replace")
                        return re.sub(r"<[^>]+>", " ", html)
                for part in payload.get("parts", []):
                    t = _extract(part)
                    if t.strip():
                        return t
                return ""

            email_dict = {
                "id": m["id"], "thread_id": full["threadId"],
                "from": sender, "subject": subject,
                "body": _extract(full["payload"]),
            }
            res = detect_resume_request(email_dict)
            if res.get("is_resume_request"):
                tag = f"RESUME REQUEST  skills={res['skills']}"
                resume_requests.append((subject, res["skills"]))
            else:
                tag = f"not a request  ({res.get('reason', '')})"
            print(f"        → {tag}\n")
        except HttpError as e:
            print(f"        [error] {e}\n")

    print(f"Resume requests found: {len(resume_requests)}")
    for subj, skills in resume_requests:
        print(f"  \"{subj[:55]}\"  skills={skills}")
except Exception as e:
    print(f"[Gmail] Skipped — {e}")

print()

# 1. Section detection
doc      = Document(RESUME)
sections = _find_experience_sections(doc)
print(f"Experience sections found: {len(sections)}")
for s in sections:
    print(f"  [{s['start_idx']:02d}-{s['end_idx']:02d}]  {s['title'][:70]}")

# 2. Bullet generation
print("\nTesting Claude bullet generation ...")
bullet = _generate_skill_bullet(
    "Toyota | Plano, TX   Jun 2025 - Present  Sr. QA Automation Engineer",
    ["Terraform", "Kubernetes"],
)
print(f"  Generated: \"{bullet}\"")

# 3. Full DOCX update
print("\nRunning update_resume(['Terraform', 'Kubernetes']) ...")
ok = update_resume(["Terraform", "Kubernetes"])
print(f"Result: {'SUCCESS' if ok else 'FAILED'}")

if ok:
    print("\nVerifying inserted paragraphs in resume_updated.docx ...")
    updated = Document("resume_updated.docx")
    for i, p in enumerate(updated.paragraphs):
        if "Terraform" in p.text or "Kubernetes" in p.text:
            print(f"  Para {i}: {p.text[:110]}")
